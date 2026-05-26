"""DOCX assembly using python-docx.

Phase 3 produces one DOCX per template_use *and* one combined package DOCX
that strings the rendered clauses together with a cover page and a table of
contents placeholder. The package step is intentionally simple — Word users
update the ToC with F9 after opening.
"""
from __future__ import annotations

from dataclasses import dataclass
from datetime import UTC, datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt


@dataclass(frozen=True)
class ClauseInput:
    clause_key: str
    title: str
    prose: str
    citations: list[str]


def render_clause_docx(out_path: Path, clause: ClauseInput, mandate_name: str) -> Path:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    _apply_default_styles(doc)
    _add_footer(doc, mandate_name, clause.clause_key)

    doc.add_heading(clause.title, level=1)
    _add_prose(doc, clause.prose)
    if clause.citations:
        doc.add_heading("Zitate", level=2)
        for c in clause.citations:
            doc.add_paragraph(c, style="List Bullet")
    doc.save(str(out_path))
    return out_path


def render_package_docx(
    out_path: Path,
    *,
    mandate_name: str,
    track: str,
    clauses: list[ClauseInput],
    version: int = 1,
) -> Path:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    _apply_default_styles(doc)
    _add_footer(doc, mandate_name, f"v{version}")

    # Cover
    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.add_run(f"MiCAR Authorization Package: {track.upper()}")
    run.bold = True
    run.font.size = Pt(18)

    sub = doc.add_paragraph(mandate_name)
    sub.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sub.runs[0].font.size = Pt(12)

    meta = doc.add_paragraph()
    meta.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    meta.add_run(
        f"Version {version} · erzeugt {datetime.now(UTC).strftime('%Y-%m-%d')}"
    ).font.size = Pt(10)

    doc.add_page_break()

    # ToC placeholder. Word fills this on F9.
    doc.add_heading("Inhaltsverzeichnis", level=1)
    doc.add_paragraph(
        "Bitte in Word: Verweise > Inhaltsverzeichnis aktualisieren (F9)."
    )
    doc.add_page_break()

    # Clauses
    for clause in clauses:
        doc.add_heading(clause.title, level=1)
        _add_prose(doc, clause.prose)
        if clause.citations:
            doc.add_heading("Zitate", level=2)
            for c in clause.citations:
                doc.add_paragraph(c, style="List Bullet")
        doc.add_page_break()

    doc.save(str(out_path))
    return out_path


def _apply_default_styles(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)


def _add_footer(doc: Document, mandate: str, suffix: str) -> None:
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.text = f"{mandate} · {suffix} · MiCAR Authorization Co-Pilot"
    p.runs[0].font.size = Pt(8)


def _add_prose(doc: Document, prose: str) -> None:
    for raw_para in prose.split("\n\n"):
        text = raw_para.strip()
        if not text:
            continue
        if text.startswith("## "):
            doc.add_heading(text[3:].strip(), level=2)
        elif text.startswith("### "):
            doc.add_heading(text[4:].strip(), level=3)
        else:
            doc.add_paragraph(text)
